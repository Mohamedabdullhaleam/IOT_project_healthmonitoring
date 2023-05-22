import schedule
import time
import pandas as pd
from docx import Document
from datetime import datetime
import telebot
import matplotlib.pyplot as plt
from docx.shared import Inches
import matplotlib.dates as mdates
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

##FOR GENERATING MANY WORDS
counter = 1


def extract_data_and_generate_report():
    global counter
    # read by default 1st sheet of an excel file
    df = pd.read_excel('data/edited_data .xlsx', sheet_name='daata')
    important_data = df[['Measurement', 'Average', 'Max', 'Min']]

    ## specify x and y  axises
    x_column = 'Time'  # Replace 'XColumn' with the name of the x-axis column
    y_column = 'BPM'  # Replace 'YColumn' with the name of the y-axis column
    df[x_column] = pd.to_datetime(df[x_column], format='%H:%M:%S')  # Adjust the format string as needed
    ## create the plot   BPM AND TIME
    plt.figure(figsize=(6, 4)) #
    fig, ax = plt.subplots()
    ax.plot(df[x_column], df[y_column])
    ax.set_xlabel('Time')
    ax.set_ylabel('BPM')
    ax.set_title('BPM')
    ##format the time b2a
    plt.gca().xaxis.set_major_formatter(mdates.DateFormatter('%H:%M:%S'))
    plt.gcf().autofmt_xdate()

    ## save the plot
    plot_filename = 'BPM_plot.png'
    plt.savefig(plot_filename)

    ## create a hisogram for SPO2
    plt.figure(figsize=(6, 4))
    plt.hist(df['SPO2'], bins=10)  # Adjust the number of bins as needed
    plt.xlabel('SPO2')
    plt.ylabel('Frequency')
    plt.title('Histogram')
    hist_filename = "histogram.png"
    plt.savefig(hist_filename)

    # Create a new Word document
    # Create a new Word document with a timestamp in the name

    doc_name = f'reports/report{counter}.docx'
    doc_name_print = f'report{counter}.docx'
    counter += 1
    doc = Document()

    # Add a heading to the document
    doc.add_heading('Important Data', level=0).alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # Convert the important data to a table in Word
    table = doc.add_table(1, len(important_data.columns))
    table.style = 'Light List Accent 1'
    new_width = Cm(10)  # Set the new width (here, 10 centimeters)
    change_table_dimensions(table, new_width)

    # # Add the header rows.
    for j in range(len(important_data.columns)):
        table.cell(0, j).text = important_data.columns[j]

    # # Add the data rows.
    for i in range(3):
        # # Add a row to the table
        row_cells = table.add_row().cells
        # # Add the data from the dataframe
        for j in range(len(important_data.columns)):
            row_cells[j].text = str(important_data.values[i, j])
    # print(important_data)

    ## plotting
    # doc.add_picture(plot_filename, width=Inches(3))
    paragraph = doc.add_paragraph()
    run = paragraph.add_run()
    run.add_picture(plot_filename, width=Inches(5))
    ## add histo'
    paragraph = doc.add_paragraph()
    his = paragraph.add_run()
    his.add_picture(hist_filename, width=Inches(5))

    print(doc_name_print + " is generated")
    # Save the Word document
    doc.save(doc_name)

    # # send word document to telegram
    bot = telebot.TeleBot("6140302269:AAG5rMISL5xamoIG5dnJcDuaJOK9qt1vWQU")
    chat_id = "639643824"
    doc = open(doc_name, 'rb')
    bot.send_document(chat_id, doc)
    print("Document sent successfully")


# Schedule the script to run every 5 minutes
schedule.every(1).seconds.do(extract_data_and_generate_report)
# adjusting table dimensions
from docx import Document
from docx.shared import Cm

def change_table_dimensions(table, width):
    # Change the table width
    table.width = width

    # Adjust the cell widths
    for row in table.rows:
        for cell in row.cells:
            cell.width = width / len(row.cells)




while True:
    schedule.run_pending()
    time.sleep(1)
