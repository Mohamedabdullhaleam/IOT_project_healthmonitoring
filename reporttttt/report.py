import schedule
import time
# import pandas lib as pd
import pandas as pd
from docx import Document
from datetime import datetime
import telebot

counter = 1
def extract_data_and_generate_report():
    global counter
    # read by default 1st sheet of an excel file
    df = pd.read_excel('data/edited_data .xlsx')

    important_data = df[['Date', 'Time', 'BPM']]

    # Create a new Word document
    # Create a new Word document with a timestamp in the name
    doc_name = f'reports/report{counter}.docx'
    doc_name_print = f'report{counter}.docx'
    counter += 1
    doc = Document()

    # Add a heading to the document
    doc.add_heading('Important Data', level=1)

    # Convert the important data to a table in Word
    table = doc.add_table(1, len(important_data.columns))
    table.style = 'Table Grid'

    # # Add the header rows.
    for j in range(len(important_data.columns)):
        table.cell(0,j).text = important_data.columns[j]


    # # Add the data rows.
    for i in range(len(important_data)):
        # # Add a row to the table
        row_cells = table.add_row().cells
        # # Add the data from the dataframe
        for j in range(len(important_data.columns)):
            row_cells[j].text = str(important_data.values[i,j])

    #print(important_data)
    print(doc_name_print + " is generated")

    # Save the Word document
    doc.save(doc_name)

    # # send word document to telegram
    bot = telebot.TeleBot("6140302269:AAG5rMISL5xamoIG5dnJcDuaJOK9qt1vWQU")
    chat_id = "639643824"
    doc = open(doc_name, 'rb')
    bot.send_document(chat_id, doc)
    print("Document sent successfully")

 # Schedule the script to run every 5 minutes
schedule.every(1).seconds.do(extract_data_and_generate_report)

while True:
    schedule.run_pending()
    time.sleep(1)